from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
OUTPUT = ROOT / "homework1_report_template.docx"


def set_run_font(run, size=11, bold=False, color=None):
    run.font.name = "Microsoft YaHei"
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)


def add_heading(document, text, level=1):
    paragraph = document.add_heading(level=level)
    run = paragraph.add_run(text)
    set_run_font(run, size=16 if level == 1 else 13, bold=True, color=(31, 78, 121))
    return paragraph


def add_placeholder(document, title):
    paragraph = document.add_paragraph()
    run = paragraph.add_run(f"[截图位置] {title}")
    set_run_font(run, size=11, bold=True, color=(192, 0, 0))
    table = document.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    cell.text = "\n\n请将对应执行结果截图粘贴到此处\n\n"
    for cell_paragraph in cell.paragraphs:
        cell_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in cell_paragraph.runs:
            set_run_font(run, size=10, color=(128, 128, 128))


def main():
    document = Document()
    section = document.sections[0]
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.4)
    section.right_margin = Cm(2.4)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Python数据分析 作业一报告")
    set_run_font(run, size=20, bold=True, color=(31, 78, 121))

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Scrapy爬取股票数据并保存到CSV与MySQL")
    set_run_font(run, size=12)

    add_heading(document, "一、作业要求", 1)
    p = document.add_paragraph()
    set_run_font(
        p.add_run(
            "使用Scrapy框架爬取股票数据，将爬取结果分别保存到CSV文件和MySQL数据库中，"
            "并把关键执行结果截图整理到Word文档中。"
        )
    )

    add_heading(document, "二、数据来源与字段说明", 1)
    p = document.add_paragraph()
    set_run_font(
        p.add_run(
            "本作业使用东方财富行情接口作为股票数据来源，爬取A股股票列表中的行情信息。"
            "主要字段包括股票代码、股票名称、最新价、涨跌幅、涨跌额、成交量、成交额、最高价、最低价和爬取时间。"
        )
    )

    table = document.add_table(rows=1, cols=3)
    table.style = "Light Shading Accent 1"
    headers = ["字段", "含义", "保存位置"]
    for index, header in enumerate(headers):
        run = table.rows[0].cells[index].paragraphs[0].add_run(header)
        set_run_font(run, bold=True)
    rows = [
        ("code", "股票代码", "CSV / MySQL"),
        ("name", "股票名称", "CSV / MySQL"),
        ("latest_price", "最新价", "CSV / MySQL"),
        ("change_percent", "涨跌幅", "CSV / MySQL"),
        ("volume", "成交量", "CSV / MySQL"),
        ("turnover", "成交额", "CSV / MySQL"),
        ("crawl_time", "爬取时间", "CSV / MySQL"),
    ]
    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            run = cells[index].paragraphs[0].add_run(value)
            set_run_font(run, size=10)

    add_heading(document, "三、项目结构", 1)
    p = document.add_paragraph()
    set_run_font(
        p.add_run(
            "项目位于 hw1 目录下，Scrapy源码保存在 src/stock_spider，CSV结果保存在 data/stocks.csv，"
            "MySQL建表脚本为 sql/create_table.sql。"
        )
    )

    add_heading(document, "四、运行过程截图", 1)
    add_placeholder(document, "依赖安装成功")
    add_placeholder(document, "MySQL数据库和数据表创建成功")
    add_placeholder(document, "Scrapy爬虫运行成功")
    add_placeholder(document, "CSV文件保存结果")
    add_placeholder(document, "MySQL查询结果")

    add_heading(document, "五、结果说明", 1)
    p = document.add_paragraph()
    set_run_font(
        p.add_run(
            "运行爬虫后，程序会自动解析股票行情JSON数据，并通过Scrapy FEEDS配置导出CSV文件。"
            "同时，MySQLPipeline会把每条股票记录写入 stock_quotes 表中。"
            "如果重复运行，程序会根据股票代码和爬取时间避免重复插入同一批次数据。"
        )
    )

    add_heading(document, "六、总结", 1)
    p = document.add_paragraph()
    set_run_font(
        p.add_run(
            "本次作业完成了Scrapy项目创建、网页数据请求、字段解析、CSV保存和MySQL持久化，"
            "体现了Python网络爬虫与数据存储的基本流程。"
        )
    )

    document.save(OUTPUT)


if __name__ == "__main__":
    main()

