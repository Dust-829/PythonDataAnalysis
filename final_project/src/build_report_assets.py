from __future__ import annotations

from pathlib import Path

import pandas as pd
from docx import Document
from docx.shared import Inches

from common import FIGURES_DIR, METRICS_DIR, PROCESSED_DIR, REPORT_DIR, TABLES_DIR, ensure_project_dirs


def add_picture_if_exists(document: Document, path: Path, width: float = 5.8) -> None:
    if path.exists():
        document.add_picture(str(path), width=Inches(width))


def add_table_from_csv(document: Document, path: Path, max_rows: int = 8) -> None:
    if not path.exists():
        document.add_paragraph(f"未生成表格：{path.name}")
        return
    df = pd.read_csv(path).head(max_rows)
    table = document.add_table(rows=1, cols=len(df.columns))
    table.style = "Table Grid"
    for i, col in enumerate(df.columns):
        table.rows[0].cells[i].text = str(col)
    for _, row in df.iterrows():
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = str(value)


def main() -> None:
    ensure_project_dirs()
    document = Document()
    document.add_heading("Python 数据分析大作业实验报告", level=0)
    document.add_paragraph("题目：电影评分与受欢迎度分析")

    document.add_heading("一、实验目的", level=1)
    document.add_paragraph("综合运用 Python 数据采集、清洗、统计分析、可视化和数据挖掘模型，对电影评分和受欢迎度进行分析，并提炼实践结论。")

    document.add_heading("二、实验环境", level=1)
    document.add_paragraph("运行环境：Python 3.x；主要库：pandas、numpy、matplotlib、seaborn、scikit-learn、plotly、streamlit。")

    document.add_heading("三、数据采集与说明", level=1)
    document.add_paragraph("数据来源包括 MovieLens 评分与标签数据、IMDb 非商业电影元数据和外部评分数据。")
    data_dict = REPORT_DIR.parent / "data" / "data_dictionary.md"
    if data_dict.exists():
        document.add_paragraph(data_dict.read_text(encoding="utf-8")[:2500])

    document.add_heading("四、数据清洗与预处理", level=1)
    document.add_paragraph("处理内容包括缺失值、重复值、异常值、数据类型转换、电影类型编码、标签文本整理和多源数据合并。")
    add_picture_if_exists(document, FIGURES_DIR / "missing_values_heatmap.png")

    document.add_heading("五、统计分析与可视化", level=1)
    document.add_paragraph("下表展示核心数值字段的描述性统计结果。")
    add_table_from_csv(document, TABLES_DIR / "descriptive_statistics.csv")
    for name in [
        "line_rating_year.png",
        "bar_top_genres.png",
        "pie_rating_distribution.png",
        "scatter_release_year_rating.png",
        "heatmap_correlation.png",
        "qq_movie_rating_mean.png",
    ]:
        add_picture_if_exists(document, FIGURES_DIR / name)

    document.add_heading("六、模型分析", level=1)
    document.add_paragraph("模型部分包括回归、分类和聚类。回归用于预测评分，分类用于识别高分评分记录，聚类用于划分电影画像。")
    for path in [METRICS_DIR / "regression_metrics.csv", METRICS_DIR / "classification_metrics.csv", METRICS_DIR / "clustering_metrics.csv"]:
        document.add_heading(path.stem, level=2)
        add_table_from_csv(document, path)
    for name in ["regression_predictions.png", "classification_roc_curves.png", "cluster_kmeans.png", "cluster_agglomerative.png"]:
        add_picture_if_exists(document, FIGURES_DIR / name)

    document.add_heading("七、创新点", level=1)
    document.add_paragraph("项目构建 Streamlit 展示页面，集中展示数据概览、可视化图表、模型指标和实践结论；同时增加标签词云作为文本分析展示。")

    document.add_heading("八、实验总结", level=1)
    document.add_paragraph("最终结论应在完整运行后结合实际图表和模型指标补充，包括热门电影类型、评分趋势、高分电影特征和聚类画像。")

    output = REPORT_DIR / "final_experiment_report.docx"
    document.save(output)
    print(f"Report generated: {output}")


if __name__ == "__main__":
    main()

