from pathlib import Path
import json

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parent.parent
REPORT_DIR = ROOT / "report"
OUTPUT_DIR = ROOT / "output"
DOC_PATH = REPORT_DIR / "homework2_report.docx"
FALLBACK_DOC_PATH = REPORT_DIR / "homework2_report_fixed.docx"
SUMMARY_JSON = OUTPUT_DIR / "analysis_summary.json"


def set_run_font(run, size=11, bold=False, color=None):
    font = run.font
    font.name = "Microsoft YaHei"
    font.size = Pt(size)
    font.bold = bold

    # Force Word to use the same font for Chinese, English, and numbers.
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.get_or_add_rFonts()
    for key in ("ascii", "hAnsi", "eastAsia", "cs"):
        r_fonts.set(qn(f"w:{key}"), "Microsoft YaHei")

    if color:
        font.color.rgb = RGBColor(*color)


def add_paragraph(document, text, size=11, bold=False):
    paragraph = document.add_paragraph()
    run = paragraph.add_run(text)
    set_run_font(run, size=size, bold=bold)
    return paragraph


def add_heading(document, text):
    paragraph = document.add_paragraph()
    run = paragraph.add_run(text)
    set_run_font(run, size=16, bold=True, color=(31, 78, 121))
    return paragraph


def add_image(document, title, path, width_cm=15.5):
    add_paragraph(document, title, size=11, bold=True)
    if path.exists():
        document.add_picture(str(path), width=Cm(width_cm))
        document.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        add_paragraph(document, f"[missing image] {path.name}")


def add_placeholder(document, title):
    add_paragraph(document, f"[截图位置] {title}", size=11, bold=True)
    table = document.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    cell.text = "\n请粘贴对应终端或运行结果截图\n"
    for paragraph in cell.paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in paragraph.runs:
            set_run_font(run, size=10, color=(128, 128, 128))


def load_summary():
    if not SUMMARY_JSON.exists():
        return None
    return json.loads(SUMMARY_JSON.read_text(encoding="utf-8"))


def build_report():
    REPORT_DIR.mkdir(parents=True, exist_ok=True)
    document = Document()

    section = document.sections[0]
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.4)
    section.right_margin = Cm(2.4)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Python数据分析 作业二报告")
    set_run_font(run, size=20, bold=True, color=(31, 78, 121))

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Pandas股票数据分析（Yahoo Finance真实数据）")
    set_run_font(run, size=12)

    summary = load_summary()

    add_heading(document, "一、作业要求与技术路线")
    add_paragraph(
        document,
        "本作业按照要求使用 yfinance 从雅虎财经下载真实股票历史数据，再使用 pandas 完成数据清洗，"
        "使用 matplotlib 和 seaborn 完成数据可视化，并进一步计算移动平均线、RSI、收益率和波动率等指标。"
    )

    add_heading(document, "二、环境准备")
    add_paragraph(
        document,
        "本项目使用的核心库包括 yfinance、pandas、matplotlib 和 seaborn。安装完成后运行分析脚本，"
        "即可自动生成原始数据、清洗结果、图表和统计摘要。"
    )
    add_placeholder(document, "依赖安装成功截图")

    add_heading(document, "三、数据获取")
    if summary:
        add_paragraph(
            document,
            f"本次分析的股票代码为 {summary['ticker']}，时间范围为 "
            f"{summary['start_date']} 至 {summary['end_date']}，共获取 {summary['row_count']} 条交易日记录。"
        )
    else:
        add_paragraph(document, "本节展示从 Yahoo Finance 获取真实股票历史数据的执行结果。")
    add_placeholder(document, "运行脚本与原始数据预览截图")
    add_image(document, "图1 数据预览表", OUTPUT_DIR / "data_preview.png")

    add_heading(document, "四、数据清洗")
    if summary:
        before_missing = sum(summary["missing_before"].values())
        after_missing = sum(summary["missing_after"].values())
        add_paragraph(
            document,
            f"首先检查缺失值，再采用前向填充与后向填充处理缺失数据。处理前缺失值总数为 {before_missing}，"
            f"处理后缺失值总数为 {after_missing}。同时删除不参与核心分析的 Adj Close 列，"
            "保留主要价格与成交量字段。"
        )
    else:
        add_paragraph(document, "本节展示缺失值检查、填充处理和无关列删除的结果。")
    add_placeholder(document, "缺失值检查与清洗结果截图")
    add_image(document, "图2 缺失值检查表", OUTPUT_DIR / "missing_values.png")

    add_heading(document, "五、数据可视化")
    add_paragraph(
        document,
        "通过收盘价折线图可以直观观察股票在研究周期内的整体趋势变化。真实数据图表由脚本自动生成并保存。"
    )
    add_image(document, "图3 收盘价走势", OUTPUT_DIR / "close_price.png")

    add_heading(document, "六、技术指标分析")
    add_paragraph(
        document,
        "移动平均线用于平滑价格波动并观察趋势方向。RSI 用于衡量短期动量，常见阈值为 70 和 30。"
    )
    add_image(document, "图4 SMA 指标图", OUTPUT_DIR / "sma.png")
    add_image(document, "图5 RSI 指标图", OUTPUT_DIR / "rsi.png")

    add_heading(document, "七、深入分析")
    add_image(document, "图6 累计收益率曲线", OUTPUT_DIR / "cumulative_return.png")
    if summary:
        add_paragraph(
            document,
            f"根据真实数据计算结果，该股票在样本期内总收益率为 {summary['total_return']}%，"
            f"日波动率为 {summary['daily_volatility']}%，年化波动率为 {summary['annualized_volatility']}%。"
            f"最新 RSI(14) 为 {summary['latest_rsi']}，最新 SMA20 为 {summary['latest_sma_20']}，"
            f"最新 SMA50 为 {summary['latest_sma_50']}。"
        )
    add_placeholder(document, "终端指标输出或 summary 文件截图")

    add_heading(document, "八、分析总结")
    if summary:
        trend_text = "上涨" if summary["total_return"] > 0 else "下跌"
        add_paragraph(
            document,
            f"从真实获取的 Yahoo Finance 数据可以看出，样本期内该股票整体呈现{trend_text}趋势。"
            "通过 pandas 的清洗与指标计算，可以把原始时间序列数据转化为更容易分析的结构；"
            "通过可视化与收益率、波动率分析，可以进一步评价股票表现和风险水平。"
        )
    else:
        add_paragraph(
            document,
            "本次作业完整展示了股票数据从获取、清洗到可视化与技术分析的流程。"
        )

    try:
        document.save(DOC_PATH)
        print(f"Saved report to {DOC_PATH}")
    except PermissionError:
        document.save(FALLBACK_DOC_PATH)
        print(f"Saved report to {FALLBACK_DOC_PATH}")


if __name__ == "__main__":
    build_report()
